# ABOUTME: Extracts text, metadata, and cover art from ePub, PDF, DOCX, and TXT files
# ABOUTME: Returns structured chapters with titles + book metadata for audiobook conversion
from __future__ import annotations

import logging
import re
import unicodedata
from dataclasses import dataclass, field
from pathlib import Path

logger = logging.getLogger("audiobook-api.extractor")


@dataclass
class Chapter:
    title: str
    text: str


@dataclass
class BookMetadata:
    title: str = "Unknown"
    author: str = "Unknown"
    language: str = ""
    publisher: str = ""
    year: str = ""
    description: str = ""


@dataclass
class ExtractionResult:
    chapters: list[Chapter] = field(default_factory=list)
    metadata: BookMetadata = field(default_factory=BookMetadata)
    cover_image: bytes | None = None


def _clean_text(text: str) -> str:
    """Normalize Unicode, collapse whitespace, strip junk."""
    text = unicodedata.normalize("NFKC", text)
    # Remove page numbers (standalone digits on a line)
    text = re.sub(r"^\s*\d+\s*$", "", text, flags=re.MULTILINE)
    # Collapse multiple blank lines
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Collapse multiple spaces
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


def extract(file_path: Path) -> ExtractionResult:
    """Extract chapters, metadata, and cover art from a file."""
    suffix = file_path.suffix.lower()
    if suffix == ".epub":
        return _extract_epub(file_path)
    elif suffix == ".pdf":
        return _extract_pdf(file_path)
    elif suffix == ".docx":
        return _extract_docx(file_path)
    elif suffix == ".txt":
        return _extract_txt(file_path)
    else:
        raise ValueError(f"Unsupported file format: {suffix}")


def _extract_epub(file_path: Path) -> ExtractionResult:
    """Extract from ePub using ebooklib + BeautifulSoup."""
    import ebooklib
    from bs4 import BeautifulSoup
    from ebooklib import epub

    book = epub.read_epub(str(file_path), options={"ignore_ncx": True})

    # Metadata
    meta = BookMetadata()
    _get = lambda ns, name: (book.get_metadata(ns, name) or [(("",), {})])[0][0]
    meta.title = _get("DC", "title") or file_path.stem
    meta.author = _get("DC", "creator") or "Unknown"
    meta.language = _get("DC", "language") or ""
    meta.publisher = _get("DC", "publisher") or ""
    date_raw = _get("DC", "date") or ""
    meta.year = date_raw[:4] if date_raw else ""
    meta.description = _get("DC", "description") or ""

    # Cover art
    cover_image = None
    # Method 1: ITEM_COVER type
    for item in book.get_items_of_type(ebooklib.ITEM_COVER):
        cover_image = item.get_content()
        break
    # Method 2: OPF metadata pointing to an image
    if cover_image is None:
        cover_meta = book.get_metadata("OPF", "cover")
        if cover_meta:
            cover_id = cover_meta[0][1].get("content", "")
            if cover_id:
                for item in book.get_items_of_type(ebooklib.ITEM_IMAGE):
                    if item.get_id() == cover_id:
                        cover_image = item.get_content()
                        break

    # Build TOC title map from NCX/nav
    toc_titles = {}
    def _walk_toc(toc_items):
        for item in toc_items:
            if isinstance(item, tuple):
                # (Section, children)
                _walk_toc(item[1])
            elif hasattr(item, "href") and hasattr(item, "title"):
                href_base = item.href.split("#")[0]
                toc_titles[href_base] = item.title
    _walk_toc(book.toc)

    # Extract chapters from spine
    chapters = []
    for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
        soup = BeautifulSoup(item.get_content(), "html.parser")
        text = soup.get_text(separator="\n")
        text = _clean_text(text)

        if not text or len(text) < 50:
            continue

        # Determine chapter title
        href = getattr(item, "file_name", "") or ""
        title = toc_titles.get(href.split("/")[-1], "")
        if not title:
            # Fallback: first h1/h2
            heading = soup.find(["h1", "h2"])
            title = heading.get_text(strip=True) if heading else ""
        if not title:
            title = f"Chapter {len(chapters) + 1}"

        chapters.append(Chapter(title=title, text=text))

    if not chapters:
        # Fallback: concatenate everything as one chapter
        all_text = []
        for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
            soup = BeautifulSoup(item.get_content(), "html.parser")
            all_text.append(soup.get_text(separator="\n"))
        combined = _clean_text("\n\n".join(all_text))
        if combined:
            chapters = [Chapter(title=meta.title, text=combined)]

    logger.info("ePub extracted: %d chapters, cover=%s", len(chapters), cover_image is not None)
    return ExtractionResult(chapters=chapters, metadata=meta, cover_image=cover_image)


def _extract_pdf(file_path: Path) -> ExtractionResult:
    """Extract from PDF using pypdf."""
    from pypdf import PdfReader

    reader = PdfReader(str(file_path))

    # Metadata
    meta = BookMetadata()
    info = reader.metadata
    if info:
        meta.title = info.title or file_path.stem
        meta.author = info.author or "Unknown"
        if info.creation_date:
            meta.year = str(info.creation_date.year)

    # Extract all page text
    pages = []
    for page in reader.pages:
        text = page.extract_text() or ""
        pages.append(text)

    full_text = _clean_text("\n\n".join(pages))

    # Try to detect chapters via heading patterns
    chapter_pattern = re.compile(
        r"^(Chapter\s+\d+[.:]*\s*.*|CHAPTER\s+\d+[.:]*\s*.*|Part\s+\d+[.:]*\s*.*|PART\s+\d+[.:]*\s*.*)$",
        re.MULTILINE,
    )
    matches = list(chapter_pattern.finditer(full_text))

    chapters = []
    if matches:
        for i, match in enumerate(matches):
            start = match.start()
            end = matches[i + 1].start() if i + 1 < len(matches) else len(full_text)
            title = match.group().strip()
            text = _clean_text(full_text[start:end])
            if text and len(text) > 50:
                chapters.append(Chapter(title=title, text=text))
    else:
        # Fall back to splitting every ~20 pages
        chunk_size = 20
        for i in range(0, len(pages), chunk_size):
            chunk_pages = pages[i : i + chunk_size]
            text = _clean_text("\n\n".join(chunk_pages))
            if text and len(text) > 50:
                chapters.append(Chapter(title=f"Section {len(chapters) + 1}", text=text))

    if not chapters and full_text:
        chapters = [Chapter(title=meta.title, text=full_text)]

    logger.info("PDF extracted: %d chapters", len(chapters))
    return ExtractionResult(chapters=chapters, metadata=meta, cover_image=None)


def _extract_docx(file_path: Path) -> ExtractionResult:
    """Extract from DOCX using python-docx."""
    from docx import Document

    doc = Document(str(file_path))

    # Metadata
    meta = BookMetadata()
    props = doc.core_properties
    meta.title = props.title or file_path.stem
    meta.author = props.author or "Unknown"

    # Split on headings
    chapters = []
    current_title = ""
    current_paragraphs: list[str] = []

    for para in doc.paragraphs:
        if para.style and para.style.name and para.style.name.startswith("Heading"):
            # Save previous chapter
            if current_paragraphs:
                text = _clean_text("\n\n".join(current_paragraphs))
                if text and len(text) > 50:
                    chapters.append(Chapter(
                        title=current_title or f"Chapter {len(chapters) + 1}",
                        text=text,
                    ))
            current_title = para.text.strip()
            current_paragraphs = []
        else:
            if para.text.strip():
                current_paragraphs.append(para.text)

    # Save last chapter
    if current_paragraphs:
        text = _clean_text("\n\n".join(current_paragraphs))
        if text and len(text) > 50:
            chapters.append(Chapter(
                title=current_title or f"Chapter {len(chapters) + 1}",
                text=text,
            ))

    if not chapters:
        all_text = _clean_text("\n\n".join(p.text for p in doc.paragraphs if p.text.strip()))
        if all_text:
            chapters = [Chapter(title=meta.title, text=all_text)]

    logger.info("DOCX extracted: %d chapters", len(chapters))
    return ExtractionResult(chapters=chapters, metadata=meta, cover_image=None)


def _extract_txt(file_path: Path) -> ExtractionResult:
    """Extract from plain text file."""
    # Try UTF-8 first, fall back to Latin-1
    try:
        text = file_path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        text = file_path.read_text(encoding="latin-1")

    text = _clean_text(text)
    meta = BookMetadata(title=file_path.stem)

    # Try to detect chapters
    chapter_pattern = re.compile(
        r"^(Chapter\s+\d+[.:]*\s*.*|CHAPTER\s+\d+[.:]*\s*.*)$",
        re.MULTILINE,
    )
    matches = list(chapter_pattern.finditer(text))

    chapters = []
    if matches:
        for i, match in enumerate(matches):
            start = match.start()
            end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
            title = match.group().strip()
            chunk = _clean_text(text[start:end])
            if chunk and len(chunk) > 50:
                chapters.append(Chapter(title=title, text=chunk))
    else:
        # Single chapter
        chapters = [Chapter(title=meta.title, text=text)]

    logger.info("TXT extracted: %d chapters", len(chapters))
    return ExtractionResult(chapters=chapters, metadata=meta, cover_image=None)
